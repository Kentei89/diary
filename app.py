import streamlit as st
from datetime import datetime, timezone, timedelta
import io
import uuid
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

KST = timezone(timedelta(hours=9))

def now_kst():
    return datetime.now(KST)

import firebase_admin
from firebase_admin import credentials, firestore

import speech_recognition as sr
try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False


st.set_page_config(
    page_title="현제 생각",
    page_icon="💭",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
    .block-container { padding-top: 1rem; padding-bottom: 4rem; max-width: 700px; }
    .stButton > button {
        border-radius: 12px; font-weight: 600;
        min-height: 48px; font-size: 1rem;
    }
    h1 { font-size: 1.8rem !important; text-align: center; }
    @media (max-width: 640px) {
        .stTextArea textarea  { font-size: 16px !important; }
        .stTextInput input    { font-size: 16px !important; }
        .stSelectbox select   { font-size: 16px !important; }
        div[data-baseweb="select"] * { font-size: 16px !important; }
        .stDateInput input    { font-size: 16px !important; }
        .stTabs [data-baseweb="tab"] { font-size: 1rem; padding: 10px 16px; }
        .block-container { padding-left: 1rem; padding-right: 1rem; }
    }
</style>
""", unsafe_allow_html=True)


# ── Firebase ──────────────────────────────────────────────────────────────────
def init_firebase():
    if firebase_admin._apps:
        return firestore.client()
    try:
        info = {
            "type":             st.secrets["firebase"]["type"],
            "project_id":       st.secrets["firebase"]["project_id"],
            "private_key_id":   st.secrets["firebase"]["private_key_id"],
            "private_key":      st.secrets["firebase"]["private_key"].replace("\\n", "\n"),
            "client_email":     st.secrets["firebase"]["client_email"],
            "client_id":        st.secrets["firebase"]["client_id"],
            "auth_uri":         "https://accounts.google.com/o/oauth2/auth",
            "token_uri":        "https://oauth2.googleapis.com/token",
        }
        cred = credentials.Certificate(info)
        firebase_admin.initialize_app(cred)
        return firestore.client()
    except Exception:
        return None


db = init_firebase()


# ── DB helpers ────────────────────────────────────────────────────────────────
def save_diary(title, content, mood, weather, date_str):
    if not db:
        return False, "Firebase 연결 실패"
    try:
        doc_id = str(uuid.uuid4())
        db.collection("diaries").document(doc_id).set({
            "id":         doc_id,
            "title":      title,
            "content":    content,
            "mood":       mood,
            "weather":    weather,
            "date":       date_str,
            "created_at": now_kst().isoformat(),
        })
        return True, None
    except Exception as e:
        return False, str(e)


def update_diary(doc_id, title, content, mood, weather, date_str):
    if not db:
        return False, "Firebase 연결 실패"
    try:
        db.collection("diaries").document(doc_id).update({
            "title":      title,
            "content":    content,
            "mood":       mood,
            "weather":    weather,
            "date":       date_str,
            "updated_at": now_kst().isoformat(),
        })
        return True, None
    except Exception as e:
        return False, str(e)


def load_diaries():
    if not db:
        return []
    try:
        docs = db.collection("diaries").order_by(
            "date", direction=firestore.Query.DESCENDING
        ).stream()
        return [doc.to_dict() for doc in docs]
    except Exception as e:
        st.error(f"목록 로드 실패: {e}")
        return []


def delete_diary(doc_id):
    if db:
        db.collection("diaries").document(doc_id).delete()


# ── STT ───────────────────────────────────────────────────────────────────────
def transcribe_audio(audio_file):
    if not PYDUB_AVAILABLE:
        return None, "pydub가 설치되지 않았습니다."
    try:
        audio_bytes = audio_file.read()
        segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
        wav_io = io.BytesIO()
        segment.export(wav_io, format="wav")
        wav_io.seek(0)

        r = sr.Recognizer()
        with sr.AudioFile(wav_io) as source:
            data = r.record(source)
        return r.recognize_google(data, language="ko-KR"), None
    except sr.UnknownValueError:
        return None, "음성을 인식하지 못했습니다. 다시 시도해주세요."
    except sr.RequestError as e:
        return None, f"Google STT 오류: {e}"
    except Exception as e:
        return None, f"변환 오류: {e}"


# ── 내보내기 ──────────────────────────────────────────────────────────────────
def export_docx(diaries, mood_emoji, weather_emoji):
    doc = Document()

    # 제목
    title_para = doc.add_heading("💭 현제 생각", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for d in diaries:
        date_raw  = d.get("date", "")[:10]
        try:
            dt = datetime.fromisoformat(date_raw)
            date_kor = f"{dt.year}년 {dt.month:02d}월 {dt.day:02d}일"
        except Exception:
            date_kor = date_raw

        mood_e    = mood_emoji.get(d.get("mood", ""), "")
        weather_e = weather_emoji.get(d.get("weather", ""), "")
        mood_label    = next((k for k, v in MOODS.items() if v == d.get("mood")), "")
        weather_label = next((k for k, v in WEATHERS.items() if v == d.get("weather")), "")

        doc.add_paragraph("─" * 40)

        # 날짜·날씨·기분 헤더
        header = doc.add_paragraph()
        run = header.add_run(f"{date_kor}  {weather_e} {weather_label}  {mood_e} {mood_label}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        # 제목
        entry_title = doc.add_paragraph()
        t_run = entry_title.add_run(d.get("title", ""))
        t_run.bold = True
        t_run.font.size = Pt(13)

        # 내용
        doc.add_paragraph(d.get("content", ""))

        # 작성 시각
        created = d.get("updated_at") or d.get("created_at", "")
        caption = doc.add_paragraph(f"작성: {created[:16]}")
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 앱 UI ─────────────────────────────────────────────────────────────────────
st.title("💭 현제 생각")

if not db:
    st.error("Firebase 설정이 필요합니다.")
    st.markdown("`.streamlit/secrets.toml` 파일을 열어 Firebase 정보를 입력해주세요.")
    st.stop()


MOODS = {
    "💪 뿌듯해":      "proud",
    "😌 평온해":      "calm",
    "🔥 의욕넘쳐":    "motivated",
    "😤 스트레스":    "stressed",
    "🥱 나른해":      "drowsy",
    "😵 정신없어":    "overwhelmed",
    "🫠 번아웃":      "burnout",
    "🥳 신나":        "excited",
    "😬 긴장돼":      "nervous",
    "🫶 감사해":      "grateful",
    "🙃 애매해":      "mixed",
    "😶 멍해":        "blank",
    "😢 속상해":      "sad",
    "😎 쿨해":        "cool",
    "🤩 설레":        "thrilled",
}
WEATHERS = {
    "☀️ 맑음":   "sunny",
    "⛅ 구름":   "cloudy",
    "🌧️ 비":    "rainy",
    "❄️ 눈":    "snowy",
    "🌫️ 흐림":  "foggy",
}
MOOD_KEYS    = list(MOODS.keys())
WEATHER_KEYS = list(WEATHERS.keys())
MOOD_EMOJI    = {v: k.split()[0] for k, v in MOODS.items()}
WEATHER_EMOJI = {v: k.split()[0] for k, v in WEATHERS.items()}

tab_write, tab_list = st.tabs(["✏️ 새 생각", "📚 목록 보기"])


# ── 새 생각 ──────────────────────────────────────────────────────────────────
with tab_write:
    if "content_area" not in st.session_state:
        st.session_state.content_area = ""

    selected_date = st.date_input("날짜", value=now_kst())

    col_mood, col_weather = st.columns(2)
    with col_mood:
        selected_mood = st.selectbox("기분", MOOD_KEYS)
    with col_weather:
        selected_weather = st.selectbox("날씨", WEATHER_KEYS)

    title = st.text_input("제목", placeholder="오늘 하루를 한 줄로 표현해보세요...")

    with st.expander("🎤 음성", expanded=False):
        if not PYDUB_AVAILABLE:
            st.warning("음성 입력을 사용하려면 `pip install pydub` 후 ffmpeg를 설치해주세요.\n\nWindows: `winget install ffmpeg`")
        else:
            audio_input = st.audio_input("녹음 버튼을 눌러 말하세요")
            if audio_input:
                if st.button("🔄 텍스트로 변환", use_container_width=True):
                    with st.spinner("음성 인식 중..."):
                        text, err = transcribe_audio(audio_input)
                    if text:
                        sep = " " if st.session_state.content_area else ""
                        st.session_state.content_area += sep + text
                        st.success(f"변환 완료: {text}")
                        st.rerun()
                    else:
                        st.error(err)

    st.text_area(
        "내용",
        placeholder="오늘 어떤 일이 있었나요? 생각이나 느낌을 자유롭게 적어보세요.",
        height=300,
        key="content_area",
        label_visibility="collapsed",
    )

    col_save, col_clear = st.columns([3, 1])
    with col_save:
        if st.button("💾 저장하기", type="primary", use_container_width=True):
            if not title.strip():
                st.error("제목을 입력해주세요.")
            elif not st.session_state.content_area.strip():
                st.error("내용을 입력해주세요.")
            else:
                ok, err = save_diary(
                    title.strip(),
                    st.session_state.content_area.strip(),
                    MOODS[selected_mood],
                    WEATHERS[selected_weather],
                    selected_date.isoformat(),
                )
                if ok:
                    st.success("저장되었습니다! 📖")
                    del st.session_state["content_area"]
                    st.rerun()
                else:
                    st.error(f"저장 실패: {err}")
    with col_clear:
        if st.button("🗑️ 지우기", use_container_width=True):
            del st.session_state["content_area"]
            st.rerun()


# ── 목록 보기 ─────────────────────────────────────────────────────────────────
with tab_list:
    diaries = load_diaries()

    if not diaries:
        st.info("📝 아직 작성된 생각이 없어요. 첫 번째 생각을 써보세요!")
    else:
        # 검색
        search = st.text_input("🔍 검색", placeholder="제목 또는 내용으로 검색...", label_visibility="collapsed")

        # 연도/월 필터
        years = sorted({d.get("date", "")[:4] for d in diaries if d.get("date")}, reverse=True)
        col_y, col_m = st.columns(2)
        with col_y:
            sel_year = st.selectbox("연도", ["전체"] + years, key="filter_year")
        with col_m:
            month_options = ["전체"] + [f"{m:02d}월" for m in range(1, 13)]
            sel_month = st.selectbox("월", month_options, key="filter_month")

        filtered = diaries
        if sel_year != "전체":
            filtered = [d for d in filtered if d.get("date", "").startswith(sel_year)]
        if sel_month != "전체":
            m = sel_month[:2]
            filtered = [d for d in filtered if d.get("date", "")[5:7] == m]
        if search.strip():
            kw = search.strip().lower()
            filtered = [d for d in filtered if kw in d.get("title", "").lower() or kw in d.get("content", "").lower()]

        col_cnt, col_exp = st.columns([2, 1])
        with col_cnt:
            st.markdown(f"총 **{len(filtered)}**개")
        with col_exp:
            docx_bytes = export_docx(filtered, MOOD_EMOJI, WEATHER_EMOJI)
            st.download_button(
                "📥 워드 내보내기",
                data=docx_bytes,
                file_name=f"생각_{now_kst().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        if "editing_id" not in st.session_state:
            st.session_state.editing_id = None

        for diary in filtered:
            doc_id    = diary["id"]
            date_str  = diary.get("date", "")[:10]
            title_d   = diary.get("title", "제목 없음")
            mood_d    = MOOD_EMOJI.get(diary.get("mood", ""), "")
            weather_d = WEATHER_EMOJI.get(diary.get("weather", ""), "")

            with st.expander(f"{weather_d} {mood_d} {date_str}  —  {title_d}"):
                if st.session_state.editing_id == doc_id:
                    # 수정 폼
                    mood_val    = next((k for k, v in MOODS.items() if v == diary.get("mood")), MOOD_KEYS[0])
                    weather_val = next((k for k, v in WEATHERS.items() if v == diary.get("weather")), WEATHER_KEYS[0])
                    try:
                        date_val = datetime.fromisoformat(diary.get("date", "")).date()
                    except Exception:
                        date_val = now_kst().date()

                    e_date    = st.date_input("날짜", value=date_val, key=f"e_date_{doc_id}")
                    e_mood    = st.selectbox("기분", MOOD_KEYS, index=MOOD_KEYS.index(mood_val), key=f"e_mood_{doc_id}")
                    e_weather = st.selectbox("날씨", WEATHER_KEYS, index=WEATHER_KEYS.index(weather_val), key=f"e_weather_{doc_id}")
                    e_title   = st.text_input("제목", value=title_d, key=f"e_title_{doc_id}")
                    e_content = st.text_area("내용", value=diary.get("content", ""), height=200, key=f"e_content_{doc_id}")

                    col_ok, col_cancel = st.columns(2)
                    with col_ok:
                        if st.button("✅ 저장", key=f"save_{doc_id}", type="primary", use_container_width=True):
                            ok, err = update_diary(doc_id, e_title.strip(), e_content.strip(), MOODS[e_mood], WEATHERS[e_weather], e_date.isoformat())
                            if ok:
                                st.session_state.editing_id = None
                                st.rerun()
                            else:
                                st.error(f"수정 실패: {err}")
                    with col_cancel:
                        if st.button("✖ 취소", key=f"cancel_{doc_id}", use_container_width=True):
                            st.session_state.editing_id = None
                            st.rerun()
                else:
                    st.write(diary.get("content", ""))
                    st.caption(f"작성: {diary.get('created_at', '')[:16]}")
                    col_edit, col_del = st.columns(2)
                    with col_edit:
                        if st.button("✏️ 수정", key=f"edit_{doc_id}", use_container_width=True):
                            st.session_state.editing_id = doc_id
                            st.rerun()
                    with col_del:
                        if st.button("🗑️ 삭제", key=f"del_{doc_id}", type="secondary", use_container_width=True):
                            delete_diary(doc_id)
                            st.rerun()
